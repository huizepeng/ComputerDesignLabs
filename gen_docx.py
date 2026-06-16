from docx import Document
from docx.shared import Pt, RGBColor
import re

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def hdr(text, level=1):
    doc.add_heading(text, level=level)

def para(text, bold=False, size=None):
    pp = doc.add_paragraph()
    r = pp.add_run(text)
    r.bold = bold
    if size: r.font.size = Pt(size)

def cblock(text):
    pp = doc.add_paragraph()
    r = pp.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

data = open(r'C:\Users\Administrator\Desktop\final_lab5\final_report.md', 'r', encoding='utf-8').read()
lines = data.split('\n')

i = 0
in_code = False
code_lines = []

while i < len(lines):
    line = lines[i]

    if line.startswith('```'):
        if in_code:
            cblock('\n'.join(code_lines))
            code_lines = []
        in_code = not in_code
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    if line.startswith('# '):
        hdr(line[2:], 1)
    elif line.startswith('## '):
        hdr(line[3:], 2)
    elif line.startswith('### '):
        hdr(line[4:], 3)
    elif line.startswith('---') or line.startswith('> '):
        pass
    elif line.startswith('|') or line.strip() == '':
        pass
    elif line.startswith('- ') or line.startswith('* '):
        para('  ' + line[2:])
    elif '**' in line:
        pp = doc.add_paragraph()
        parts = re.split(r'(\*\*[^*]+\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                r = pp.add_run(part[2:-2])
                r.bold = True
            else:
                pp.add_run(part)
    elif '`' in line:
        pp = doc.add_paragraph()
        parts = re.split(r'(\\`[^`]+\\`)', line)
        for part in parts:
            if part.startswith('`') and part.endswith('`'):
                r = pp.add_run(part[1:-1])
                r.font.name = 'Consolas'
                r.font.size = Pt(9)
            else:
                pp.add_run(part)
    else:
        para(line)

    i += 1

doc.save(r'C:\Users\Administrator\Desktop\final_lab5\final_report.docx')
print('OK: final_report.docx')
